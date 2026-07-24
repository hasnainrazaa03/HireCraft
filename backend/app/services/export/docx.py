"""Render a Master Resume to a .docx file.

A clean, single-column, ATS-friendly Word document — the format recruiters and
some portals still ask for. Layout mirrors the LaTeX résumé's structure so the
two exports read consistently. python-docx has no HTML/markup escaping concern;
values are written as literal runs.
"""

from __future__ import annotations

import io

from app.schemas.resume import MasterResume

_ACCENT = "1A3E6F"
_MUTED = "555555"

# Month map for "2024-05" -> "May 2024" date display.
_MONTHS = [
    "", "Jan", "Feb", "Mar", "apr", "May", "Jun",
    "Jul", "Aug", "Sep", "Oct", "Nov", "Dec",
]


def _fmt_date(value: str | None) -> str:
    if not value:
        return ""
    if value == "Present":
        return "Present"
    if "-" in value:
        year, month = value.split("-", 1)
        idx = int(month) if month.isdigit() else 0
        return f"{_MONTHS[idx].title()} {year}" if 0 < idx < 13 else year
    return value


def _date_range(start: str | None, end: str | None) -> str:
    left, right = _fmt_date(start), _fmt_date(end)
    return f"{left} – {right}" if left and right else (left or right)


def resume_to_docx(resume: MasterResume) -> bytes:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = docx.Document()

    # Tighter default margins than Word's 1 inch.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(46)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    b = resume.basics

    # --- Header ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(b.name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    if b.headline:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(b.headline)
        run.font.color.rgb = RGBColor.from_string(_MUTED)

    contact_bits = [
        x
        for x in (b.location, b.phone, str(b.email))
        if x
    ]
    for url in (b.linkedin, b.github, b.website):
        if url:
            contact_bits.append(str(url).split("://")[-1].rstrip("/"))
    if contact_bits:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run("  •  ".join(contact_bits)).font.size = Pt(9.5)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(_ACCENT)
        # A bottom border approximates the LaTeX section rule.
        _bottom_border(p)

    def entry(title: str, right: str, subtitle: str = "", sub_right: str = "") -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        tab = _right_tab(p)
        r = p.add_run(title)
        r.bold = True
        if right:
            p.add_run(f"\t{right}")
        if subtitle or sub_right:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(1)
            _right_tab(sp)
            si = sp.add_run(subtitle)
            si.italic = True
            si.font.color.rgb = RGBColor.from_string(_MUTED)
            if sub_right:
                sr = sp.add_run(f"\t{sub_right}")
                sr.italic = True
                sr.font.color.rgb = RGBColor.from_string(_MUTED)
        _ = tab

    def bullets(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    sections = resume.section_order
    for section in sections:
        if section == "summary" and (b.summary or b.headline):
            if b.summary:
                heading("Summary")
                doc.add_paragraph(b.summary)
        elif section == "education" and resume.education:
            heading("Education")
            for e in resume.education:
                degree = f"{e.degree}"
                if e.field_of_study:
                    degree += f", {e.field_of_study}"
                if e.gpa:
                    degree += f" — GPA {e.gpa}"
                entry(e.institution, _date_range(e.start_date, e.end_date), degree, e.location or "")
                if e.coursework:
                    doc.add_paragraph(f"Relevant Coursework: {', '.join(e.coursework)}")
                bullets(e.highlights)
        elif section == "experience" and resume.experience:
            heading("Experience")
            for x in resume.experience:
                entry(x.company, _date_range(x.start_date, x.end_date), x.title, x.location or "")
                bullets(x.highlights)
                if x.technologies:
                    tp = doc.add_paragraph()
                    tr = tp.add_run(f"Technologies: {', '.join(x.technologies)}")
                    tr.italic = True
                    tr.font.size = Pt(9.5)
        elif section == "projects" and resume.projects:
            heading("Projects")
            for p in resume.projects:
                entry(p.name, _date_range(p.start_date, p.end_date), p.description or "")
                bullets(p.highlights)
                if p.technologies:
                    tp = doc.add_paragraph()
                    tr = tp.add_run(f"Built with: {', '.join(p.technologies)}")
                    tr.italic = True
                    tr.font.size = Pt(9.5)
        elif section == "skills" and resume.skills:
            heading("Skills")
            for g in resume.skills:
                sp = doc.add_paragraph()
                sp.add_run(f"{g.category}: ").bold = True
                sp.add_run(", ".join(g.items))
        elif section == "certifications" and resume.certifications:
            heading("Certifications")
            bullets([
                c.name + (f" — {c.issuer}" if c.issuer else "") + (f" ({c.date})" if c.date else "")
                for c in resume.certifications
            ])
        elif section == "awards" and resume.awards:
            heading("Awards")
            bullets([
                a.title + (f", {a.issuer}" if a.issuer else "") + (f" ({a.date})" if a.date else "")
                for a in resume.awards
            ])
        elif section == "publications" and resume.publications:
            heading("Publications")
            bullets([
                (", ".join(pub.authors) + ". " if pub.authors else "")
                + pub.title
                + (f", {pub.venue}" if pub.venue else "")
                + (f", {pub.date}" if pub.date else "")
                for pub in resume.publications
            ])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def cover_letter_to_docx(
    resume: MasterResume,
    paragraphs: list[str],
    *,
    company: str | None = None,
    role: str | None = None,
    hiring_manager: str | None = None,
    date_line: str | None = None,
) -> bytes:
    """Render a cover letter to a clean, single-column .docx.

    Mirrors the LaTeX cover-letter layout: sender header, date, recipient block,
    salutation, body paragraphs, and a sign-off — so the DOCX and PDF read the
    same.
    """
    import docx
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(54)
        section.left_margin = section.right_margin = Pt(60)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.15

    b = resume.basics

    # Sender header.
    name = doc.add_paragraph()
    run = name.add_run(b.name)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(_ACCENT)
    name.paragraph_format.space_after = Pt(2)

    contact_bits = [str(b.email)]
    for value in (b.phone, b.location):
        if value:
            contact_bits.append(str(value))
    contact = doc.add_paragraph()
    contact_run = contact.add_run("  ·  ".join(contact_bits))
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor.from_string(_MUTED)
    contact.paragraph_format.space_after = Pt(14)

    if date_line:
        doc.add_paragraph(date_line)

    # Recipient block.
    recipient_lines = [
        line
        for line in (hiring_manager, company, role and f"Re: {role}")
        if line
    ]
    for line in recipient_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
    if recipient_lines:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_paragraph(f"Dear {hiring_manager or 'Hiring Manager'},")

    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

    closing = doc.add_paragraph("Sincerely,")
    closing.paragraph_format.space_after = Pt(2)
    signature = doc.add_paragraph()
    sig_run = signature.add_run(b.name)
    sig_run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _right_tab(paragraph) -> None:
    """Add a right-aligned tab stop at the text margin so a trailing '\\t' pushes
    the date to the right edge."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(7.0), WD_TAB_ALIGNMENT.RIGHT
    )


def _bottom_border(paragraph) -> None:
    """Draw a thin bottom border under a paragraph (the section rule)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _ACCENT)
    borders.append(bottom)
    p_pr.append(borders)
