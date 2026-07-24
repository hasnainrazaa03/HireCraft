"""Extract plain text from an uploaded résumé file.

Deterministic, per-format text extraction — no LLM here. The cleaned text is
handed to ``structure.py``, which asks Gemini to map it onto the Master Resume
schema. Keeping extraction separate means a parsing failure is attributable to a
specific stage, and the text step stays fast and free.
"""

from __future__ import annotations

import io
import re

from app.core.logging import get_logger

logger = get_logger(__name__)

MAX_TEXT_CHARS = 40_000


class ExtractionError(Exception):
    """Raised when a file's text cannot be extracted."""


def _clean(text: str) -> str:
    text = text.replace("\x00", "")
    # Collapse runs of blank lines and trailing whitespace.
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_TEXT_CHARS]


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("PDF support is not installed.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # Many résumé PDFs are "encrypted" with an empty owner password.
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ExtractionError(
                    "This PDF is password-protected. Remove the password and try again."
                ) from exc
        pages = [page.extract_text() or "" for page in reader.pages]
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Could not read that PDF: {exc}") from exc

    text = _clean("\n".join(pages))
    if len(text) < 40:
        raise ExtractionError(
            "Almost no text could be read from that PDF — it may be a scanned image. "
            "Try a text-based PDF, or paste your résumé instead."
        )
    return text


def extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support is not installed.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read that Word document: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs]
    # Tables are common in résumés (skills grids, two-column layouts).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = _clean("\n".join(parts))
    if len(text) < 40:
        raise ExtractionError("That Word document appears to be empty.")
    return text


# LaTeX control sequences and structural noise we strip before handing to the LLM.
_TEX_COMMENT = re.compile(r"(?<!\\)%.*$", re.MULTILINE)
_TEX_COMMAND = re.compile(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?")
_TEX_BRACES = re.compile(r"[{}]")


def extract_latex(data: bytes) -> str:
    raw = data.decode("utf-8", errors="replace")
    text = _TEX_COMMENT.sub("", raw)
    # Common escaped characters back to their literal form.
    for escaped, literal in (
        (r"\&", "&"), (r"\%", "%"), (r"\$", "$"), (r"\#", "#"),
        (r"\_", "_"), (r"\{", "{"), (r"\}", "}"),
        (r"\textbackslash", "\\"), (r"\\", "\n"), (r"~", " "),
    ):
        text = text.replace(escaped, literal)
    text = _TEX_COMMAND.sub(" ", text)
    text = _TEX_BRACES.sub(" ", text)
    text = _clean(text)
    if len(text) < 40:
        raise ExtractionError("Could not extract text from that LaTeX file.")
    return text


def extract_text(data: bytes) -> str:
    text = _clean(data.decode("utf-8", errors="replace"))
    if len(text) < 40:
        raise ExtractionError("That file is too short to be a résumé.")
    return text


# Maps a lower-cased extension to its extractor.
_EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "tex": extract_latex,
    "latex": extract_latex,
    "txt": extract_text,
    "md": extract_text,
    "json": extract_text,  # handled specially upstream, but a safe fallback
}

SUPPORTED_EXTENSIONS = tuple(sorted(_EXTRACTORS))


def extract(filename: str, data: bytes) -> str:
    """Extract text from ``data`` based on ``filename``'s extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise ExtractionError(
            f"Unsupported file type {ext or '(none)'!r}. "
            f"Upload a PDF, Word, LaTeX, or text résumé."
        )
    text = extractor(data)
    logger.info("parsing.extracted", ext=ext, chars=len(text))
    return text
